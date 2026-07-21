"""Word builder (python-docx)."""
from __future__ import annotations

import io

from app.document_platform.generation.builders.base import (
    FIXED_BUILD_TIME,
    AbstractFileBuilder,
    BuildError,
)
from app.document_platform.generation.content_model import ContentModel


class WordBuilder(AbstractFileBuilder):
    format_name = "word"
    extension = "docx"
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    features = frozenset({"sections", "tables", "styling"})

    def build(self, model: ContentModel) -> bytes:
        try:
            from docx import Document

            doc = Document()
            props = doc.core_properties
            props.author = "UnityWorks Generation Platform"
            props.created = FIXED_BUILD_TIME.replace(tzinfo=None)
            props.modified = FIXED_BUILD_TIME.replace(tzinfo=None)

            doc.add_heading(model.title, level=0)
            if model.subtitle:
                doc.add_paragraph(model.subtitle, style="Subtitle")
            for key, value in model.metadata.items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(value)

            for section in model.sections:
                doc.add_heading(section.heading, level=max(1, min(3, section.level)))
                for para in section.paragraphs:
                    doc.add_paragraph(para)
                for bullet in section.bullets:
                    doc.add_paragraph(bullet, style="List Bullet")
                if section.table is not None:
                    table = doc.add_table(
                        rows=1, cols=max(1, len(section.table.headers)),
                    )
                    table.style = "Light Grid Accent 1"
                    for i, header in enumerate(section.table.headers):
                        cell = table.rows[0].cells[i]
                        cell.text = header
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                    for data_row in section.table.rows:
                        cells = table.add_row().cells
                        for i, value in enumerate(data_row[:len(cells)]):
                            cells[i].text = value

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except BuildError:
            raise
        except Exception as e:
            raise BuildError(f"Word build failed: {e}") from e
