"""
Document text extraction.

Supported types:
  - PDF (.pdf)            → pypdf
  - Word (.docx)          → python-docx (paragraphs + tables)
  - Plain text families   → direct decode (.txt, .md, .csv, .json, code files, …)

Legacy binary .doc is NOT supported — users get a clear error asking
for .docx or PDF instead.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Optional


class DocumentExtractionError(Exception):
    """Raised when a document cannot be parsed or its type is unsupported."""


# Extensions treated as plain text (decoded directly, no parser needed)
TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".jsonl",
    ".yaml", ".yml", ".toml", ".xml", ".html", ".htm", ".log", ".ini", ".cfg",
    ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".cs", ".go", ".rs",
    ".c", ".cpp", ".h", ".hpp", ".rb", ".php", ".sql", ".sh", ".ps1", ".css",
}

SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx"}


def is_supported_document(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


@dataclass
class ExtractedText:
    """Result of a text extraction pass."""
    text: str
    page_count: Optional[int] = None  # PDFs only

    @property
    def char_count(self) -> int:
        return len(self.text)


def extract_text(file_bytes: bytes, filename: str) -> ExtractedText:
    """
    Extract plain text from an uploaded document.

    Raises DocumentExtractionError for unsupported types or parse failures.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)
    if ext == ".docx":
        return _extract_docx(file_bytes, filename)
    if ext == ".doc":
        raise DocumentExtractionError(
            f"Legacy .doc format is not supported ({filename}). "
            "Please save the file as .docx or PDF and upload again."
        )
    if ext in TEXT_EXTENSIONS:
        return _extract_plain_text(file_bytes, filename)

    raise DocumentExtractionError(
        f"Unsupported document type: {ext or 'no extension'} ({filename}). "
        f"Supported: PDF, Word (.docx), and text files."
    )


def _extract_pdf(file_bytes: bytes, filename: str) -> ExtractedText:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise DocumentExtractionError(
            "PDF support is not installed. Run: pip install pypdf"
        ) from e

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        if reader.is_encrypted:
            # Try empty-password decryption (common for "protected" PDFs)
            try:
                reader.decrypt("")
            except Exception:
                raise DocumentExtractionError(
                    f"PDF is password-protected: {filename}"
                )
        pages = []
        for i, page in enumerate(reader.pages):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                pages.append(f"[Page {i + 1}]\n{page_text}")
        text = "\n\n".join(pages)
        if not text.strip():
            raise DocumentExtractionError(
                f"No extractable text found in PDF: {filename}. "
                "It may be a scanned/image-only PDF — try uploading page screenshots "
                "as images instead so the vision model can read them."
            )
        return ExtractedText(text=text, page_count=len(reader.pages))
    except DocumentExtractionError:
        raise
    except Exception as e:
        raise DocumentExtractionError(f"Failed to parse PDF {filename}: {e}") from e


def _extract_docx(file_bytes: bytes, filename: str) -> ExtractedText:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise DocumentExtractionError(
            "Word support is not installed. Run: pip install python-docx"
        ) from e

    try:
        document = docx.Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables — render as pipe-separated rows so the LLM can read them
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        text = "\n\n".join(parts)
        if not text.strip():
            raise DocumentExtractionError(f"No extractable text found in Word file: {filename}")
        return ExtractedText(text=text)
    except DocumentExtractionError:
        raise
    except Exception as e:
        raise DocumentExtractionError(f"Failed to parse Word file {filename}: {e}") from e


def _extract_plain_text(file_bytes: bytes, filename: str) -> ExtractedText:
    if b"\x00" in file_bytes[:8192]:
        raise DocumentExtractionError(f"File appears to be binary, not text: {filename}")
    text = file_bytes.decode("utf-8", errors="replace").strip()
    if not text:
        raise DocumentExtractionError(f"File is empty: {filename}")
    return ExtractedText(text=text)
