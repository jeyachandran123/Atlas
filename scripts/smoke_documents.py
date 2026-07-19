"""One-shot smoke test for the documents feature. Run from backend/: python scripts/smoke_documents.py"""
import asyncio
import io
import tempfile
from pathlib import Path

import app.api.v1.chat.router as _router  # noqa: F401 — catches wiring errors
from app.documents.extractor import DocumentExtractionError, extract_text, is_supported_document
from app.db.models import MessageDocument  # noqa: F401

# Plain text
r = extract_text(b"Invoice #123\nTotal: $450.00\nDue: 2026-08-01", "invoice.txt")
assert "Total: $450.00" in r.text and r.page_count is None, r

# DOCX (built in memory)
import docx

d = docx.Document()
d.add_paragraph("Quarterly Report")
d.add_paragraph("Revenue grew 14% year over year.")
t = d.add_table(rows=2, cols=2)
t.cell(0, 0).text = "Region"
t.cell(0, 1).text = "Sales"
t.cell(1, 0).text = "EMEA"
t.cell(1, 1).text = "42000"
buf = io.BytesIO()
d.save(buf)
r2 = extract_text(buf.getvalue(), "report.docx")
assert "Revenue grew 14%" in r2.text and "EMEA | 42000" in r2.text, r2.text

# Blank PDF → clear error
from pypdf import PdfWriter

w = PdfWriter()
w.add_blank_page(width=200, height=200)
pbuf = io.BytesIO()
w.write(pbuf)
try:
    extract_text(pbuf.getvalue(), "blank.pdf")
    raise SystemExit("expected DocumentExtractionError for blank pdf")
except DocumentExtractionError as e:
    assert "No extractable text" in str(e), e

# Supported / unsupported / legacy .doc
assert is_supported_document("a.pdf") and is_supported_document("b.docx") and is_supported_document("c.md")
assert not is_supported_document("d.exe")
try:
    extract_text(b"x", "old.doc")
    raise SystemExit("expected .doc rejection")
except DocumentExtractionError as e:
    assert ".docx or PDF" in str(e), e

# Storage + context + prompt block (memory fallback, no Redis needed)
from app.documents.context import DocumentContext
from app.documents.service import DocumentService
from app.documents.storage import DocumentStorage


async def main() -> None:
    from app.storage.local import LocalBlobStorage
    tmp = Path(tempfile.mkdtemp())
    svc = DocumentService(
        storage=DocumentStorage(blobs=LocalBlobStorage(tmp)), context=DocumentContext()
    )
    att = await svc.process_upload(
        b"Employee handbook. Vacation: 20 days/year.", "handbook.txt", "text/plain", "conv-1"
    )
    assert att.char_count > 0
    assert await svc.has_documents("conv-1")
    block = await svc.build_document_block("conv-1")
    assert "handbook.txt" in block and "Vacation: 20 days/year" in block, block
    sysp, userp = await svc.build_document_prompt(
        "How many vacation days?", "conv-1", "BASE", [{"role": "user", "content": "hi"}]
    )
    assert "Document Analysis" in sysp and "BASE" in sysp
    assert "Uploaded Documents" in userp and "Current Question" in userp
    print("SMOKE_OK")


asyncio.run(main())
