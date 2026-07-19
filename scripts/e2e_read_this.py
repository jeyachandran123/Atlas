"""Regression test: vague 'Read this' prompt must produce a summary, not a refusal."""
import io
import json
import uuid

import httpx

BASE = "http://localhost:8000/api/v1"
EMAIL = f"e2e-{uuid.uuid4().hex[:8]}@example.com"
PASSWORD = "e2e-password-123"

with httpx.Client(timeout=300) as c:
    c.post(f"{BASE}/auth/register", json={
        "email": EMAIL, "password": PASSWORD, "role": "developer", "org_id": "default",
    })
    token = c.post(f"{BASE}/auth/login", json={"email": EMAIL, "password": PASSWORD}).json()["access_token"]
    H = {"Authorization": f"Bearer {token}"}

    import docx
    d = docx.Document()
    d.add_paragraph("Atlas Platform Vision")
    d.add_paragraph("Atlas is an AI coding assistant with semantic code search across large codebases.")
    d.add_paragraph("Roadmap: IDE extensions in Q3, team collaboration in Q4, enterprise SSO in 2027.")
    d.add_paragraph("Pricing: free tier for individuals, $12/user/month for teams.")
    buf = io.BytesIO()
    d.save(buf)

    files = {"documents": ("vision.docx", buf.getvalue(),
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    with c.stream("POST", f"{BASE}/chat/stream/vision", headers=H,
                  data={"message": "Read this", "agent_mode": "auto"}, files=files) as resp:
        assert resp.status_code == 200, resp.read()[:300]
        text = ""
        for line in resp.iter_lines():
            if line.startswith("data:"):
                try:
                    ev = json.loads(line[5:].strip())
                except json.JSONDecodeError:
                    continue
                if ev.get("type") == "token":
                    text += ev.get("content", "")
                if ev.get("type") == "error":
                    raise SystemExit(f"stream error: {ev}")

    print("ANSWER:", text[:400].replace("\n", " "))
    lowered = text.lower()
    refusals = ["cannot access", "can't access", "cannot read", "can't read", "unable to read", "unable to access"]
    assert not any(r in lowered for r in refusals), "MODEL REFUSED AGAIN"
    assert any(k in lowered for k in ["atlas", "roadmap", "pricing", "coding assistant"]), "summary lacks doc content"
    print("READ_THIS_OK")
