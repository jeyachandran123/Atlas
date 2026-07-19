"""End-to-end test of chat + document upload against a running API on :8000."""
import io
import json
import sys
import uuid

import httpx

BASE = "http://localhost:8000/api/v1"
EMAIL = f"e2e-{uuid.uuid4().hex[:8]}@example.com"
PASSWORD = "e2e-password-123"


def sse_events(resp: httpx.Response) -> list[dict]:
    events = []
    for line in resp.iter_lines():
        if line.startswith("data:"):
            try:
                events.append(json.loads(line[5:].strip()))
            except json.JSONDecodeError:
                pass
    return events


def summarize(events: list[dict]) -> tuple[str, dict | None, dict | None]:
    text = "".join(e.get("content", "") for e in events if e.get("type") == "token")
    done = next((e for e in events if e.get("type") == "done"), None)
    err = next((e for e in events if e.get("type") == "error"), None)
    return text, done, err


with httpx.Client(timeout=300) as c:
    # 0. Ensure default org exists (mirror of firebase auto-registration)
    #    Register may 500 if org missing — create via SQL is out of scope here;
    #    org 'default' is created by firebase logins; try register and report.
    r = c.post(f"{BASE}/auth/register", json={
        "email": EMAIL, "password": PASSWORD, "full_name": "E2E Test",
        "role": "developer", "org_id": "default",
    })
    print("register:", r.status_code, r.text[:200])
    if r.status_code not in (200, 201):
        sys.exit("REGISTER_FAILED")

    r = c.post(f"{BASE}/auth/login", json={"email": EMAIL, "password": PASSWORD})
    print("login:", r.status_code)
    assert r.status_code == 200, r.text
    token = r.json()["access_token"]
    H = {"Authorization": f"Bearer {token}"}

    # 1. Plain text chat stream
    with c.stream("POST", f"{BASE}/chat/stream", headers=H,
                  json={"message": "Reply with exactly: HELLO_OK", "agent_mode": "auto"}) as resp:
        print("stream status:", resp.status_code)
        assert resp.status_code == 200, resp.read()[:300]
        events = sse_events(resp)
    text, done, err = summarize(events)
    print("stream text:", text[:120].replace("\n", " "))
    print("stream done:", bool(done), "| error:", err)
    assert done and not err, f"text stream failed: {err}"

    # 2. Upload a Word file to a NEW conversation
    import docx
    d = docx.Document()
    d.add_paragraph("Project Phoenix Budget")
    d.add_paragraph("The total approved budget is 78500 dollars.")
    d.add_paragraph("Project lead: Meridian Chen. Deadline: 30 November 2026.")
    buf = io.BytesIO()
    d.save(buf)

    files = {"documents": ("phoenix.docx", buf.getvalue(),
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    data = {"message": "What is the total approved budget?", "agent_mode": "auto"}
    with c.stream("POST", f"{BASE}/chat/stream/vision", headers=H, data=data, files=files) as resp:
        print("doc upload status:", resp.status_code)
        assert resp.status_code == 200, resp.read()[:500]
        events = sse_events(resp)
    text, done, err = summarize(events)
    print("doc answer:", text[:200].replace("\n", " "))
    print("doc done:", bool(done), "| error:", err)
    assert done and not err, f"doc stream failed: {err}"
    conv_id = done["conversation_id"]

    # 3. Follow-up question WITHOUT re-upload (plain JSON /chat/stream)
    with c.stream("POST", f"{BASE}/chat/stream", headers=H,
                  json={"message": "Who is the project lead?", "conversation_id": conv_id}) as resp:
        print("follow-up status:", resp.status_code)
        assert resp.status_code == 200, resp.read()[:300]
        events = sse_events(resp)
    text, done, err = summarize(events)
    print("follow-up answer:", text[:200].replace("\n", " "))
    assert done and not err, f"follow-up failed: {err}"

    # 4. Messages list includes the document attachment
    r = c.get(f"{BASE}/chat/conversations/{conv_id}/messages", headers=H)
    print("messages status:", r.status_code)
    assert r.status_code == 200, r.text[:300]
    msgs = r.json()
    doc_msgs = [m for m in msgs if m.get("documents")]
    print("messages:", len(msgs), "| with documents:", len(doc_msgs))
    assert doc_msgs, "no message carries the document attachment"
    doc = doc_msgs[0]["documents"][0]
    print("attachment:", doc["filename"], doc["url"])

    # 5. Download the document back
    r = c.get(f"{BASE}{doc['url'].replace('/chat', '/chat', 1)}", headers=H)
    r2 = c.get(f"{BASE}{doc['url']}", headers=H)
    print("download status:", r2.status_code, "bytes:", len(r2.content))
    assert r2.status_code == 200 and len(r2.content) > 0

    print("E2E_OK")
